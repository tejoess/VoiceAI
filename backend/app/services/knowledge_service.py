"""Knowledge document ingestion service.

Handles file parsing (PDF / DOCX / TXT), token-aware chunking, and async
indexing into ChromaDB.  Each chunk carries a stable deterministic id so
re-uploading the same doc is idempotent (ChromaDB upserts).

Chunking strategy
-----------------
* Split on double-newlines (paragraph boundaries) first.
* If a paragraph exceeds _CHUNK_TOKENS, split further on single newlines
  then on spaces (word-boundary fallback).
* 400-token target chunk size with a 50-token overlap between adjacent
  chunks so semantic meaning doesn't get cut mid-sentence.
"""

from __future__ import annotations

import asyncio
import hashlib
import io
import uuid
from pathlib import Path

from sqlalchemy.ext.asyncio import AsyncSession

from app.core.logging import get_logger
from app.models.knowledge import KnowledgeDocument
from app.voice.providers.rag import add_chunks, delete_document_chunks

log = get_logger(__name__)

_CHUNK_TOKENS = 400
_OVERLAP_TOKENS = 50


# ── Token counting ─────────────────────────────────────────────────────────

def _tokenizer():
    try:
        import tiktoken
        return tiktoken.get_encoding("cl100k_base")
    except Exception:  # noqa: BLE001
        return None


_ENC = None


def _count_tokens(text: str) -> int:
    global _ENC
    if _ENC is None:
        _ENC = _tokenizer()
    if _ENC is None:
        return len(text) // 4  # rough fallback: ~4 chars/token
    return len(_ENC.encode(text))


# ── Text extraction ────────────────────────────────────────────────────────

def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def _extract_text(data: bytes, file_type: str) -> str:
    if file_type == "pdf":
        return _extract_pdf(data)
    if file_type == "docx":
        return _extract_docx(data)
    return _extract_txt(data)


# ── Chunking ───────────────────────────────────────────────────────────────

def _split_into_chunks(text: str) -> list[str]:
    """Split text into token-bounded chunks with overlap."""
    # Paragraph-level split first
    paragraphs: list[str] = [p.strip() for p in text.split("\n\n") if p.strip()]

    chunks: list[str] = []
    current: list[str] = []
    current_tokens = 0

    for para in paragraphs:
        p_tokens = _count_tokens(para)

        # Paragraph fits in remaining space — accumulate it.
        if current_tokens + p_tokens <= _CHUNK_TOKENS:
            current.append(para)
            current_tokens += p_tokens
        else:
            # Flush what we have so far.
            if current:
                chunks.append("\n\n".join(current))

            # If the paragraph itself is larger than the chunk limit, split it.
            if p_tokens > _CHUNK_TOKENS:
                words = para.split()
                sub: list[str] = []
                sub_tokens = 0
                for word in words:
                    wt = _count_tokens(word + " ")
                    if sub_tokens + wt > _CHUNK_TOKENS and sub:
                        chunks.append(" ".join(sub))
                        # keep overlap: last N tokens worth of words
                        sub = sub[-_OVERLAP_TOKENS:]
                        sub_tokens = _count_tokens(" ".join(sub))
                    sub.append(word)
                    sub_tokens += wt
                if sub:
                    chunks.append(" ".join(sub))
                current = []
                current_tokens = 0
            else:
                # Start fresh with this paragraph.
                current = [para]
                current_tokens = p_tokens

    if current:
        chunks.append("\n\n".join(current))

    return [c for c in chunks if c.strip()]


def _make_chunk_id(doc_id: str, index: int) -> str:
    return f"{doc_id}-chunk-{index:04d}"


# ── Public API ─────────────────────────────────────────────────────────────

ALLOWED_TYPES: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    # browsers sometimes send these
    "application/msword": "docx",
    "application/octet-stream": "txt",
}

ALLOWED_EXTENSIONS: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
    ".md": "txt",
}


def detect_file_type(filename: str, content_type: str) -> str | None:
    ext = Path(filename).suffix.lower()
    if ext in ALLOWED_EXTENSIONS:
        return ALLOWED_EXTENSIONS[ext]
    return ALLOWED_TYPES.get(content_type)


async def ingest_document(
    db: AsyncSession,
    agent_id: str,
    filename: str,
    file_type: str,
    data: bytes,
) -> KnowledgeDocument:
    """Parse, chunk, embed, and index a document.  Returns the DB record."""
    doc = KnowledgeDocument(
        agent_id=uuid.UUID(agent_id),
        filename=filename,
        file_type=file_type,
        status="processing",
        doc_metadata={"size_bytes": len(data)},
    )
    db.add(doc)
    await db.flush()
    doc_id = str(doc.id)

    # Run the CPU-bound extraction and embedding in a background thread
    # so the event loop stays free for incoming audio frames.
    async def _run_in_thread():
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, _extract_text, data, file_type)
        chunks_text = await loop.run_in_executor(None, _split_into_chunks, text)
        chunk_dicts = [
            {
                "id": _make_chunk_id(doc_id, i),
                "text": chunk,
                "metadata": {
                    "doc_id": doc_id,
                    "agent_id": agent_id,
                    "filename": filename,
                    "chunk_index": i,
                },
            }
            for i, chunk in enumerate(chunks_text)
        ]
        await add_chunks(agent_id, chunk_dicts)
        return len(chunks_text)

    try:
        chunk_count = await _run_in_thread()
        doc.status = "ready"
        doc.chunk_count = chunk_count
        log.info("knowledge.ingested", doc_id=doc_id, chunks=chunk_count, file=filename)
    except Exception as exc:  # noqa: BLE001
        doc.status = "error"
        doc.error_message = str(exc)[:500]
        log.error("knowledge.ingest_failed", doc_id=doc_id, error=str(exc))

    await db.commit()
    await db.refresh(doc)
    return doc


async def delete_document(db: AsyncSession, agent_id: str, doc_id: str) -> bool:
    """Delete a document's DB record and its ChromaDB chunks."""
    result = await db.get(KnowledgeDocument, uuid.UUID(doc_id))
    if result is None or str(result.agent_id) != agent_id:
        return False
    await db.delete(result)
    await delete_document_chunks(agent_id, doc_id)
    await db.commit()
    return True
